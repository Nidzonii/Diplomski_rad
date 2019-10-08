import re
import os
import docx
from recnik import *
from spellchecker import SpellChecker
import platform

def funkcija_jedan(imeFajla, izlazniFajl):
    izmenjeniPodaci = list()
    pomocniBrojac = 1
    duzinaReci = 0
    if izlazniFajl != imeFajla:
            izlazniFajl = izlazniFajl + '.' + tip_fajla(imeFajla)
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        izmenjeniPodaci.append(podaci[0].upper())
        while pomocniBrojac < len(podaci):
            if podaci[pomocniBrojac] == '.':
                izmenjeniPodaci.append(podaci[pomocniBrojac])
                duzinaReci = pomocniBrojac - 1
                brojIzaTacke = izmenjeniPodaci[duzinaReci]
                while podaci[duzinaReci].isalpha():
                    duzinaReci -= 1
                duzinaReci += 1
                if ''.join(podaci[duzinaReci:pomocniBrojac]) in skracenice or brojIzaTacke.isnumeric():
                    pomocniBrojac += 1    
                else:
                    try:
                        pomocniBrojac += 1
                        while not podaci[pomocniBrojac].isalpha():
                            izmenjeniPodaci.append(podaci[pomocniBrojac])
                            pomocniBrojac += 1
                        if podaci[pomocniBrojac].isupper() == False:
                            izmenjeniPodaci.append(podaci[pomocniBrojac].upper())
                            pomocniBrojac += 1
                    except IndexError:
                        pass
            else:
                izmenjeniPodaci.append(podaci[pomocniBrojac])
                pomocniBrojac += 1
        izmenjeniPodaci = ''.join(izmenjeniPodaci)
        fajl.close()
        izmenjeniFajl = open(izlazniFajl, 'w', encoding='utf8')
        izmenjeniFajl.write(izmenjeniPodaci)
        izmenjeniFajl.close()
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        noviFajl = docx.Document()
        proveraNarednogParagrafa = False
        spajanjeIstihRunova = False
        brojac = 0
        stiloviParagrafa = list()
        for stil in range(len(fajl.paragraphs)):
            stiloviParagrafa.append(fajl.paragraphs[stil].style)
        for para in fajl.paragraphs:
            paragraf = noviFajl.add_paragraph()
            prviRun = True
            for run in para.runs:
                listaKaraktera = list(run.text)
                if listaKaraktera and prviRun:
                    listaKaraktera[0] = listaKaraktera[0].upper()
                    prviRun = False
                if proveraNarednogParagrafa:
                    prvoSlovo = 0
                    try:
                        while not listaKaraktera[prvoSlovo].isalpha():
                            prvoSlovo += 1
                        listaKaraktera[prvoSlovo] = listaKaraktera[prvoSlovo].upper()
                    except:
                        pass
                    proveraNarednogParagrafa = False
                for karakter in range(len(listaKaraktera)-1):
                    if listaKaraktera[karakter] == '.':
                        duzinaReci = karakter - 1
                        brojIzaTacke = listaKaraktera[duzinaReci]
                        if type(listaKaraktera[duzinaReci]) == int:
                            continue
                        while listaKaraktera[duzinaReci].isalpha():
                            duzinaReci -= 1
                        duzinaReci += 1
                        if ''.join(listaKaraktera[duzinaReci:karakter]) in skracenice or brojIzaTacke.isnumeric():
                            continue
                        else:
                            try:
                                duzinaReci = karakter + 1
                                while not listaKaraktera[duzinaReci].isalpha():
                                    duzinaReci += 1
                                listaKaraktera[duzinaReci] = listaKaraktera[duzinaReci].upper()
                            except IndexError:
                                pass
                if listaKaraktera:
                    try:
                        if listaKaraktera[-1] == '.' and not listaKaraktera[-2].isnumeric():
                            proveraNarednogParagrafa = True
                    except:
                        pass
                    else:
                        pocetnaPozicija = len(listaKaraktera) - 1
                        try:
                            while not listaKaraktera[pocetnaPozicija].isalnum() and listaKaraktera[pocetnaPozicija] != '.':
                                pocetnaPozicija -= 1
                            if listaKaraktera[pocetnaPozicija] == '.' and not listaKaraktera[pocetnaPozicija - 1].isnumeric():
                                proveraNarednogParagrafa = True
                        except:
                            pass
                r = svojstva_runa(paragraf, run, ''.join(listaKaraktera))
            paragraf.style = stiloviParagrafa[brojac]
            paragraf.alignment = para.alignment
            brojac += 1
        noviFajl.save(izlazniFajl)
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')

def funkcija_dva(imeFajla, izlazniFajl):
    listaKaraktera = list()
    if izlazniFajl != imeFajla:
            izlazniFajl = izlazniFajl + '.' + tip_fajla(imeFajla)
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        for karakter in podaci:
            listaKaraktera.append(karakter)
        for i in range(len(listaKaraktera)-1):
            if (listaKaraktera[i] == ',' or listaKaraktera[i] == '.') and not listaKaraktera[i+1].isspace():
                listaKaraktera.insert(i+1, ' ')
        izmenjeniPodaci = ''.join(listaKaraktera)
        fajl.close()
        izmenjeniFajl = open(izlazniFajl, 'w', encoding='utf8')
        izmenjeniFajl.write(izmenjeniPodaci)
        izmenjeniFajl.close()
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        noviFajl = docx.Document()
        brojac = 0
        stiloviParagrafa = list()
        for stil in range(len(fajl.paragraphs)):
            stiloviParagrafa.append(fajl.paragraphs[stil].style)
        for para in fajl.paragraphs:
            paragraf = noviFajl.add_paragraph()
            trenutniRun = 0
            for run in para.runs:
                listaKaraktera = list(run.text)
                try:
                    if listaKaraktera and listaKaraktera[-1] == '.' or listaKaraktera[-1] == ',':
                        prviKarakterSlRun = fajl.paragraphs[brojac].runs[trenutniRun+1].text[0]
                        if not prviKarakterSlRun.isspace():
                            listaKaraktera.append(' ')
                except:
                    pass  
                trenutniRun += 1
                for i in range(len(listaKaraktera)-1):
                    if (listaKaraktera[i] == ',' or listaKaraktera[i] == '.') and not listaKaraktera[i+1].isspace():
                        listaKaraktera.insert(i+1, ' ')
                r = svojstva_runa(paragraf, run, ''.join(listaKaraktera))
            paragraf.style = stiloviParagrafa[brojac]
            paragraf.alignment = para.alignment
            brojac += 1
        noviFajl.save(izlazniFajl)
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')
        
def funkcija_tri(imeFajla, izlazniFajl):
    if izlazniFajl != imeFajla:
            izlazniFajl = izlazniFajl + '.' + tip_fajla(imeFajla)
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        izmenjeniPodaci = re.sub(' +', ' ', podaci)
        fajl.close()
        izmenjeniFajl = open(izlazniFajl, 'w', encoding='utf8')
        izmenjeniFajl.write(izmenjeniPodaci)
        izmenjeniFajl.close()
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        noviFajl = docx.Document()
        brojac = 0
        stiloviParagrafa = list()
        for stil in range(len(fajl.paragraphs)):
            stiloviParagrafa.append(fajl.paragraphs[stil].style)
        for para in fajl.paragraphs:
            paragraf = noviFajl.add_paragraph()
            for run in para.runs:
                izmenjeniPodaci = re.sub(' +', ' ', run.text)
                r = svojstva_runa(paragraf, run, izmenjeniPodaci)
            paragraf.style = stiloviParagrafa[brojac]
            paragraf.alignment = para.alignment
            brojac += 1
        noviFajl.save(izlazniFajl)
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')   
        
def funkcija_cetiri(imeFajla, izbor, izlazniFajl):
    recnikReci = ucitajRecnik()
    listaKaraktera = list()
    novaLista = list()
    rec = list()
    invertovanaSlova = invertovanjeRecnika(slova)
    if izlazniFajl != imeFajla:
            izlazniFajl = izlazniFajl + '.' + tip_fajla(imeFajla)
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        for karakter in podaci:
            listaKaraktera.append(karakter)
        if izbor.upper() == 'Ć' or izbor.upper() == 'ĆIRILICA':
            for karakter in range(len(listaKaraktera)):
                if listaKaraktera[karakter].isalpha():
                    rec.append(listaKaraktera[karakter])
                else:
                    novaLista.append(''.join(rec))
                    novaLista.append(listaKaraktera[karakter])
                    rec = list()
            rec = list()
            for deo in range(len(novaLista)):
                if novaLista[deo].isalpha() and novaLista[deo][0] in invertovanaSlova:
                    continue
                elif novaLista[deo].isalpha() and (novaLista[deo].lower() in recnikReci or novaLista[deo].capitalize() in recnikReci):
                    try:
                        pomocnaLista = recnikReci[novaLista[deo].lower()]
                        novaLista[deo] = slova[novaLista[deo][0]] + pomocnaLista[1:]
                    except:
                        pass
                    try:
                        pomocnaLista = recnikReci[novaLista[deo].capitalize()]
                        novaLista[deo] = slova[novaLista[deo][0]] + pomocnaLista[1:]
                    except:
                        pass
                elif novaLista[deo].isalpha() and not novaLista[deo].lower() in recnikReci:
                    for karakter in range(len(novaLista[deo])-1):
                        if ((novaLista[deo][karakter].lower() == 'n' or novaLista[deo][karakter].lower() == 'l') and novaLista[deo][karakter+1] == 'j') or (novaLista[deo][karakter].lower() == 'd' and novaLista[deo][karakter+1] == 'ž'):
                            s = novaLista[deo][karakter] + novaLista[deo][karakter+1]
                            rec.append(slova[s])
                        elif (novaLista[deo][karakter] == 'j' and (novaLista[deo][karakter-1].lower() == 'n' or novaLista[deo][karakter-1].lower() == 'l')) or (novaLista[deo][karakter] == 'ž' and novaLista[deo][karakter-1].lower() == 'd'):
                            continue
                        else:
                            rec.append(slova[novaLista[deo][karakter]])
                    rec.append(slova[novaLista[deo][-1]])
                    novaLista[deo] = ''.join(rec)
                    rec = list()            
        elif izbor.upper() == 'L' or izbor.upper() == 'LATINICA':
            for karakter in range(len(listaKaraktera)):
                if listaKaraktera[karakter] in slova:
                    continue
                if listaKaraktera[karakter].isalpha():
                    listaKaraktera[karakter] = invertovanaSlova[listaKaraktera[karakter]]
        fajl.close()
        if izbor.upper() == 'L' or izbor.upper() == 'LATINICA':
            izmenjeniPodaci = ''.join(listaKaraktera)
        elif izbor.upper() == 'Ć' or izbor.upper() == 'ĆIRILICA':
            izmenjeniPodaci = ''.join(novaLista)
        izmenjeniFajl = open(izlazniFajl, 'w', encoding='utf8')
        izmenjeniFajl.write(izmenjeniPodaci)
        izmenjeniFajl.close()
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        noviFajl = docx.Document()
        brojac = 0
        spajanjeIstihRunova = False
        stiloviParagrafa = list()
        tekstIstogRuna = list()
        for stil in range(len(fajl.paragraphs)):
            stiloviParagrafa.append(fajl.paragraphs[stil].style)
        if izbor.upper() == 'Ć' or izbor.upper() == 'ĆIRILICA':
            for para in fajl.paragraphs:
                paragraf = noviFajl.add_paragraph()
                trenutniRun = 0 
                for run in para.runs:
                    listaKaraktera = list(run.text)
                    if listaKaraktera and listaKaraktera[-1].isalnum():
                            try:
                                prviKarakterSlRun = fajl.paragraphs[brojac].runs[trenutniRun+1].text[0]
                                if prviKarakterSlRun.isalpha():
                                    spajanjeIstihRunova = True
                                    tekstIstogRuna = tekstIstogRuna + listaKaraktera.copy()
                                    trenutniRun += 1
                                    continue
                            except:
                                pass
                    if spajanjeIstihRunova:
                        listaKaraktera = tekstIstogRuna + listaKaraktera
                        tekstIstogRuna = list()
                        spajanjeIstihRunova = False
                    trenutniRun += 1
                    for karakter in range(len(listaKaraktera)):
                        if listaKaraktera[karakter].isalpha():
                            rec.append(listaKaraktera[karakter])
                        else:
                            novaLista.append(''.join(rec))
                            novaLista.append(listaKaraktera[karakter])
                            rec = list()
                    novaLista.append(''.join(rec))
                    listaKaraktera = list()
                    rec = list()
                    for deo in range(len(novaLista)):
                        if novaLista[deo].isalpha() and novaLista[deo][0] in invertovanaSlova:
                            continue
                        elif novaLista[deo].isalpha() and (novaLista[deo].lower() in recnikReci or novaLista[deo].capitalize() in recnikReci):
                            try:
                                pomocnaLista = recnikReci[novaLista[deo].lower()]
                                novaLista[deo] = slova[novaLista[deo][0]] + pomocnaLista[1:]
                            except:
                                pass
                            try:
                                pomocnaLista = recnikReci[novaLista[deo].capitalize()]
                                novaLista[deo] = slova[novaLista[deo][0]] + pomocnaLista[1:]
                            except:
                                pass
                        elif novaLista[deo].isalpha() and not novaLista[deo].lower() in recnikReci:
                            for karakter in range(len(novaLista[deo])-1):
                                if ((novaLista[deo][karakter].lower() == 'n' or novaLista[deo][karakter].lower() == 'l') and novaLista[deo][karakter+1] == 'j') or (novaLista[deo][karakter].lower() == 'd' and novaLista[deo][karakter+1] == 'ž'):
                                    s = novaLista[deo][karakter] + novaLista[deo][karakter+1]
                                    try:
                                        rec.append(slova[s])
                                    except:
                                        pass
                                elif (novaLista[deo][karakter] == 'j' and (novaLista[deo][karakter-1].lower() == 'n' or novaLista[deo][karakter-1].lower() == 'l')) or (novaLista[deo][karakter] == 'ž' and novaLista[deo][karakter-1].lower() == 'd'):
                                    continue
                                else:
                                    try:
                                        rec.append(slova[novaLista[deo][karakter]])
                                    except:
                                        pass
                            try:
                                rec.append(slova[novaLista[deo][-1]])
                            except:
                                pass
                            novaLista[deo] = ''.join(rec)
                            rec = list()
                    #r = paragraf.add_run(''.join(novaLista))
                    r = svojstva_runa(paragraf, run, ''.join(novaLista))
                    novaLista = list()
                paragraf.style = stiloviParagrafa[brojac]
                paragraf.alignment = para.alignment
                brojac += 1
        elif izbor.upper() == 'L' or izbor.upper() == 'LATINICA':
            for para in fajl.paragraphs:
                paragraf = noviFajl.add_paragraph()
                for run in para.runs:
                    listaKaraktera = list(run.text)
                    for karakter in range(len(listaKaraktera)):
                        if listaKaraktera[karakter].isalpha() and listaKaraktera[karakter] in slova:
                            continue
                        elif listaKaraktera[karakter].isalpha() and listaKaraktera[karakter] in invertovanaSlova:
                            listaKaraktera[karakter] = invertovanaSlova[listaKaraktera[karakter]]
                    r = svojstva_runa(paragraf, run, ''.join(listaKaraktera))
                paragraf.style = stiloviParagrafa[brojac]
                paragraf.alignment = para.alignment
                brojac += 1
        noviFajl.save(izlazniFajl)
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')

def funkcija_pet(imeFajla):
    if tip_fajla(imeFajla) == 'txt':
        komanda = 'paps ' + imeFajla + ' > ' + naziv_fajla(imeFajla) + '.ps && ps2pdf ' + naziv_fajla(imeFajla) + '.ps'
        os.system(komanda)
        komanda = 'rm ' + naziv_fajla(imeFajla) + '.ps'
        os.system(komanda)
    elif tip_fajla(imeFajla) == 'docx':
        if platform.system().lower() == 'linux':
            folder = os.getcwd()
            komanda = 'lowriter --convert-to pdf --outdir "' + folder + '" ' + imeFajla
            os.system(komanda)
        elif platform.system().lower() == 'windows':
            komanda = 'officetopdf.exe ' + imeFajla + ' ' + naziv_fajla(imeFajla) + '.pdf'
            os.system(komanda)
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')

def funkcija_sest(argumenti):
    for argument in range(len(argumenti)-1):
        if tip_fajla(argumenti[argument]) != 'pdf':
            print('Morate uneti .pdf fajl!')
            return
    if tip_fajla(argumenti[-1]) != 'pdf':
        print('Izlazni fajl mora biti ekstenzije .pdf!')
        return
    if platform.system().lower() == 'linux':
        komanda = "gs -q -dBATCH -dNOPAUSE -sDEVICE=pdfwrite -sOutputFile=" + argumenti[-1] 
    elif platform.system().lower() == 'windows':
        if platform.architecture()[0] == '64bit':
            komanda = 'C:\\"Program Files"\\gs\\gs9.27\\bin\\gswin64c.exe -q -dBATCH -dNOPAUSE -sDEVICE=pdfwrite -sOutputFile=' + argumenti[-1]
        else:
            komanda = 'C:\\"Program Files"\\gs\\gs9.27\\bin\\gswin32c.exe -q -dBATCH -dNOPAUSE -sDEVICE=pdfwrite -sOutputFile=' + argumenti[-1]
    karakteri = list(komanda)
    for argument in argumenti:
        karakteri.append(' '+argument)
    komanda = ''.join(karakteri)
    os.system(komanda)

def funkcija_sedam(imeFajla, distanca, izlazniFajl):
    komanda = 'python3 diplomski_rad.py -f ' + imeFajla + ' -2'
    os.system(komanda)
    spel = SpellChecker(language=None, distance=distanca)
    spel.word_frequency.load_text_file('./sr_full.txt')
    pogresneReci = list()
    recnik = dict()
    if izlazniFajl != imeFajla:
            izlazniFajl = izlazniFajl + '.' + tip_fajla(imeFajla)
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        fajl.close()
        reci = podaci.split()
        for r in range(len(reci)):
            reci[r] = re.sub('\W+', '', reci[r])
        pogresneReci = list(spel.unknown(reci))
        pogresneReci = list(filter(None, pogresneReci))
        for pogresnaRec in pogresneReci:
            print('Reč ' + pogresnaRec +' je možda pogrešno napisana. Želite li da reč "' + pogresnaRec+ '" zamenite nekom drugom sekvencom karaktera?')
            odgovor = input('da|ne|izlaz: ')
            while odgovor != 'ne' and odgovor != 'da' and odgovor != 'izlaz':
                odgovor = input('da|ne|izlaz: ') 
            if odgovor == 'ne':
                continue
            elif odgovor == 'da':
                if len(pogresnaRec) > 10:
                    spel.distance = 1
                else:
                    spel.distance = distanca
                recnik[pogresnaRec] = list(spel.candidates(pogresnaRec))
                indeksiPreporuka = [recnik[pogresnaRec].index(x)+1 for x in recnik[pogresnaRec]]
                indeksiIReci = list()
                for i in range(len(recnik[pogresnaRec])):
                    indeksiIReci.append(recnik[pogresnaRec][i] + ' -> indeks: ' + str(indeksiPreporuka[i]))
                print('Preporuke za reč ' + pogresnaRec + ': ' + str(indeksiIReci))
                odgovor = int(input('Unesite indeks: '))
                while type(odgovor) != int or odgovor > indeksiPreporuka[-1] or odgovor < 1:
                    print('Potrebno je uneti realan broj koji je manji ili jednak broju ponuđenih reči u listi a veci od nule!')
                    odgovor = int(input('Unesite indeks: '))
                podaci = re.sub('\\b' + pogresnaRec + '\\b', recnik[pogresnaRec][odgovor-1], podaci, flags=re.I)
            elif odgovor == 'izlaz':
                izmenjeniFajl = open(imeFajla, 'w', encoding='utf8')
                izmenjeniFajl.write(podaci)
                izmenjeniFajl.close()
                return
        izmenjeniFajl = open(izlazniFajl, 'w')
        izmenjeniFajl.write(podaci)
        izmenjeniFajl.close()
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        noviFajl = docx.Document()
        podaci = list()
        brojac = 0
        stiloviParagrafa = list()
        tekstIstogRuna = list()
        spajanjeIstihRunova = False
        proveraNarednogParagrafa = False
        for stil in range(len(fajl.paragraphs)):
            stiloviParagrafa.append(fajl.paragraphs[stil].style)
        for para in fajl.paragraphs:
            for karakter in para.text:
                podaci.append(karakter)
            podaci.append('\n')
        podaci = ''.join(podaci)
        reci = podaci.split()
        for r in range(len(reci)):
            reci[r] = re.sub('\W+', '', reci[r])
        pogresneReci = list(spel.unknown(reci))
        pogresneReci = list(filter(None, pogresneReci))
        parReciZaKorekciju = dict()
        brojIspitanihReci = len(pogresneReci)
        baremJednaIspitanaRec = False
        for pogresnaRec in pogresneReci:
            print('Reč ' + pogresnaRec +' je možda pogrešno napisana. Želite li da reč "' + pogresnaRec+ '" zamenite nekom drugom sekvencom karaktera?')
            odgovor = input('da|ne|izlaz: ')
            while odgovor != 'ne' and odgovor != 'da' and odgovor != 'izlaz':
                odgovor = input('da|ne|izlaz: ') 
            if odgovor == 'ne':
                brojIspitanihReci -= 1
                continue
            elif odgovor == 'da':
                brojIspitanihReci -= 1
                if len(pogresnaRec) > 10:
                    spel.distance = 1
                else:
                    spel.distance = distanca
                recnik[pogresnaRec] = list(spel.candidates(pogresnaRec))
                indeksiPreporuka = [recnik[pogresnaRec].index(x)+1 for x in recnik[pogresnaRec]]
                indeksiIReci = list()
                for i in range(len(recnik[pogresnaRec])):
                    indeksiIReci.append(recnik[pogresnaRec][i] + ' -> indeks: ' + str(indeksiPreporuka[i]))
                print('Preporuke za reč ' + pogresnaRec + ': ' + str(indeksiIReci))
                odgovor = int(input('Unesite indeks: '))
                while type(odgovor) != int or odgovor > indeksiPreporuka[-1] or odgovor < 1:
                    print('Potrebno je uneti realan broj koji je manji ili jednak broju ponuđenih reči u listi a veci od nule!')
                    odgovor = int(input('Unesite indeks: '))
                parReciZaKorekciju[pogresnaRec] = recnik[pogresnaRec][odgovor-1]
                baremJednaIspitanaRec = True
            elif odgovor == 'izlaz' and baremJednaIspitanaRec:
                for para in fajl.paragraphs:
                    paragraf = noviFajl.add_paragraph()
                    trenutniRun = 0
                    for run in para.runs:
                        listaKaraktera = list(run.text)
                        if listaKaraktera and listaKaraktera[-1].isalnum():
                            try:
                                prviKarakterSlRun = fajl.paragraphs[brojac].runs[trenutniRun+1].text[0]
                                if prviKarakterSlRun.isalpha():
                                    spajanjeIstihRunova = True
                                    tekstIstogRuna = tekstIstogRuna + listaKaraktera.copy()
                                    trenutniRun += 1
                                    continue
                            except:
                                pass  
                        if spajanjeIstihRunova:
                            listaKaraktera = tekstIstogRuna + listaKaraktera
                            tekstIstogRuna = list()
                            spajanjeIstihRunova = False
                        trenutniRun += 1
                        tekst = ''.join(listaKaraktera)
                        for rec in parReciZaKorekciju.keys():
                            pronadjeno = re.search('\\b' + rec + '\\b', tekst, flags=re.I)
                            if pronadjeno:
                                tekst = re.sub('\\b' + rec + '\\b', parReciZaKorekciju[rec], tekst, flags=re.I)
                        r = svojstva_runa(paragraf, run, tekst)
                        trenutniRun += 1
                    paragraf.style = stiloviParagrafa[brojac]
                    paragraf.alignment = para.alignment
                    brojac += 1
                noviFajl.save(izlazniFajl)
                komanda = 'python3 diplomski_rad.py -f ' + imeFajla + ' -1'
                os.system(komanda)
                return
            elif odgovor == 'izlaz' and not baremJednaIspitanaRec:
                return
        if brojIspitanihReci < 1:
            for para in fajl.paragraphs:
                paragraf = noviFajl.add_paragraph()
                trenutniRun = 0
                for run in para.runs:
                    listaKaraktera = list(run.text)
                    if listaKaraktera and listaKaraktera[-1].isalnum():
                        try:
                            prviKarakterSlRun = fajl.paragraphs[brojac].runs[trenutniRun+1].text[0]
                            if prviKarakterSlRun.isalpha():
                                spajanjeIstihRunova = True
                                tekstIstogRuna = tekstIstogRuna + listaKaraktera.copy()
                                continue
                        except:
                            pass  
                    if spajanjeIstihRunova:
                        listaKaraktera = tekstIstogRuna + listaKaraktera
                        tekstIstogRuna = list()
                        spajanjeIstihRunova = False
                    trenutniRun += 1
                    tekst = ''.join(listaKaraktera)
                    for rec in parReciZaKorekciju.keys():
                        pronadjeno = re.search('\\b' + rec + '\\b', tekst, flags=re.I)
                        if pronadjeno:
                            tekst = re.sub('\\b' + rec + '\\b', parReciZaKorekciju[rec], tekst, flags=re.I)
                    r = svojstva_runa(paragraf, run, tekst)
                    trenutniRun += 1
                paragraf.style = stiloviParagrafa[brojac]
                paragraf.alignment = para.alignment
                brojac += 1
            noviFajl.save(izlazniFajl)
            komanda = 'python3 diplomski_rad.py -f ' + imeFajla + ' -1'
            os.system(komanda)
            return
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')
        return
        
def funkcija_osam(imeFajla, brojReci, fleg):
    komanda = 'python3 diplomski_rad.py -f ' + imeFajla + ' -2'
    os.system(komanda)
    brojPojavljivanja = 0
    brojPonavljanjaReci = dict()
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        fajl.close()
        reci = podaci.split()
        for r in range(len(reci)):
            reci[r] = re.sub('\W+', '', reci[r])
        if fleg.upper() == 'DA':
            for r in reci:
                if r.lower() in brojPonavljanjaReci:
                    brojPonavljanjaReci[r.lower()] += 1
                else:
                    brojPonavljanjaReci[r.lower()] = 1
        else:
            for r in reci:
                if r in brojPonavljanjaReci:
                    brojPonavljanjaReci[r] += 1
                else:
                    brojPonavljanjaReci[r] = 1
        najcesceReci = sorted(brojPonavljanjaReci, key=brojPonavljanjaReci.get, reverse=True)
        if brojReci > len(najcesceReci):
            brojReci = len(najcesceReci)
            print('Uneti broj je veći od broja pronađenih reči. Sledi prikazan svih pronađenih reči:')
        else:
            print('Prvih ' + str(brojReci) + ' najčešće korišćenih reči u tekstu:')
        for i in range(brojReci):
            print(str(i+1) + '. ' + najcesceReci[i] + ' -> broj ponavljanja: ' + str(brojPonavljanjaReci[najcesceReci[i]]))  
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        reci = list()
        for para in fajl.paragraphs:
            reciParagrafa = para.text.split()
            for r in range(len(reciParagrafa)):
                reciParagrafa[r] = re.sub('\W+', '', reciParagrafa[r])
                reci.append(reciParagrafa[r])
        if fleg.upper() == 'DA':
            for r in reci:
                if r.lower() in brojPonavljanjaReci:
                    brojPonavljanjaReci[r.lower()] += 1
                else:
                    brojPonavljanjaReci[r.lower()] = 1
        else:
            for r in reci:
                if r in brojPonavljanjaReci:
                    brojPonavljanjaReci[r] += 1
                else:
                    brojPonavljanjaReci[r] = 1
        najcesceReci = sorted(brojPonavljanjaReci, key=brojPonavljanjaReci.get, reverse=True)
        if brojReci > len(najcesceReci):
            brojReci = len(najcesceReci)
            print('Uneti broj je veći od broja pronađenih reči. Sledi prikazan svih pronađenih reči:')
        else:
            print('Prvih ' + str(brojReci) + ' najčešće korišćenih reči u tekstu:')
        for i in range(brojReci):
            print(str(i+1) + '. ' + najcesceReci[i] + ' -> broj ponavljanja: ' + str(brojPonavljanjaReci[najcesceReci[i]]))    
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')

def funkcija_devet(imeFajla, argumenti, fleg_zameni_sve, fleg_ignorisi, izlazniFajl):
    sekvenca = argumenti[0]
    novaSekvenca = argumenti[1]
    if izlazniFajl != imeFajla:
            izlazniFajl = izlazniFajl + '.' + tip_fajla(imeFajla)
    if tip_fajla(imeFajla) == 'txt':
        try:
            fajl = open(imeFajla, 'r', encoding='utf8')
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        podaci = fajl.read()
        pronadjeno = re.search('\\b' + sekvenca + '\\b', podaci)
        pronadjeno_i = re.search('\\b' + sekvenca + '\\b', podaci, flags=re.I)
        if fleg_zameni_sve.upper() == 'SVAKO' and fleg_ignorisi.upper() == 'DA' and pronadjeno_i:
            podaci = re.sub('\\b'+sekvenca+'\\b', novaSekvenca, podaci, flags=re.I)
        elif fleg_zameni_sve.upper() == 'PRVO' and fleg_ignorisi.upper() == 'DA' and pronadjeno_i:
            podaci = re.sub('\\b'+sekvenca+'\\b', novaSekvenca, podaci, 1, flags=re.I)
        elif fleg_zameni_sve.upper() == 'SVAKO' and fleg_ignorisi.upper() == 'NE' and pronadjeno:
            podaci = re.sub('\\b'+sekvenca+'\\b', novaSekvenca, podaci)
        elif fleg_zameni_sve.upper() == 'PRVO' and fleg_ignorisi.upper() == 'NE' and pronadjeno:
            podaci = re.sub('\\b'+sekvenca+'\\b', novaSekvenca, podaci, 1)
        fajl.close()
        izmenjeniFajl = open(izlazniFajl, 'w', encoding='utf8')
        izmenjeniFajl.write(podaci)
        izmenjeniFajl.close()
    elif tip_fajla(imeFajla) == 'docx':
        try:
            fajl = docx.Document(imeFajla)
        except:
            print('Uneli ste nepostojeći ili nedozvoljeni fajl!')
            return
        noviFajl = docx.Document()
        brojac = 0
        spajanjeIstihRunova = False
        stiloviParagrafa = list()
        tekstIstogRuna = list()
        flegDva = True
        for stil in range(len(fajl.paragraphs)):
            stiloviParagrafa.append(fajl.paragraphs[stil].style)
        for para in fajl.paragraphs:
            paragraf = noviFajl.add_paragraph()
            trenutniRun = 0
            for run in para.runs:
                listaKaraktera = list(run.text)
                if listaKaraktera and listaKaraktera[-1].isalpha():
                    try:
                        prviKarakterSlRun = fajl.paragraphs[brojac].runs[trenutniRun+1].text[0]
                        if prviKarakterSlRun.isalpha():
                            spajanjeIstihRunova = True
                            tekstIstogRuna = tekstIstogRuna + listaKaraktera.copy()
                            trenutniRun += 1
                            continue
                    except:
                        pass  
                if spajanjeIstihRunova:
                    listaKaraktera = tekstIstogRuna + listaKaraktera
                    tekstIstogRuna = list()
                    spajanjeIstihRunova = False
                trenutniRun += 1
                podaci = ''.join(listaKaraktera)
                pronadjeno = re.search('\\b' + sekvenca + '\\b', podaci)
                pronadjeno_i = re.search('\\b' + sekvenca + '\\b', podaci, flags=re.I)
                if fleg_zameni_sve.upper() == 'SVAKO' and fleg_ignorisi.upper() == 'DA' and pronadjeno_i:
                    podaci = re.sub('\\b' + sekvenca + '\\b', novaSekvenca, podaci, flags=re.I)
                elif fleg_zameni_sve.upper() == 'PRVO' and fleg_ignorisi.upper() == 'DA' and pronadjeno_i:
                    if flegDva:
                        podaci = re.sub('\\b'+sekvenca+'\\b', novaSekvenca, podaci, 1, flags=re.I)
                        flegDva = False
                elif fleg_zameni_sve.upper() == 'SVAKO' and fleg_ignorisi.upper() == 'NE' and pronadjeno:
                    podaci = re.sub('\\b' + sekvenca + '\\b', novaSekvenca, podaci)
                elif fleg_zameni_sve.upper() == 'PRVO' and fleg_ignorisi.upper() == 'NE' and pronadjeno:
                    if flegDva:
                        podaci = re.sub('\\b' + sekvenca + '\\b', novaSekvenca, podaci, 1)
                        flegDva = False
                r = svojstva_runa(paragraf, run, podaci)
            paragraf.style = stiloviParagrafa[brojac]
            paragraf.alignment = para.alignment
            brojac += 1
        noviFajl.save(izlazniFajl)    
    else:
        print('Funkcija prima kao ulaz fajlove sa ekstenzijom .txt i .docx')

def svojstva_runa(paragraf, run, tekst):
    r = paragraf.add_run(tekst)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    return r

def tip_fajla(imeFajla):
    listaTacaka = list()
    ekstenzija = list()
    for i in range(len(imeFajla)):
        if imeFajla[i] == '.':
            listaTacaka.append(i)
    for i in range(listaTacaka[-1]+1, len(imeFajla)):
        ekstenzija.append(imeFajla[i])
    return ''.join(ekstenzija)

def naziv_fajla(imeFajla):
    listaTacaka = list()
    naziv = list()
    for i in range(len(imeFajla)):
        if imeFajla[i] == '.':
            listaTacaka.append(i)
    for i in range(listaTacaka[-1]):
        naziv.append(imeFajla[i])
    return ''.join(naziv)
